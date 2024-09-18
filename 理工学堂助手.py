import io
import os
import threading
import time
import tkinter as tk
from tkinter import font
from tkinter.constants import *
import requests
import ttkbootstrap as ttk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from ttkbootstrap.dialogs import Messagebox
import keyring

base_url = "http://lgxt.wutp.com.cn/api"
headers = {
    'Accept': '*/*',
    'Accept-Language': 'zh-CN,zh;q=0.9',
    'Content-Type': 'application/x-www-form-urlencoded',
}

session = requests.Session()

SERVICE_NAME = '理工学堂'


def login(username, password):
    login_url = f"{base_url}/login"
    data = {
        'loginName': username,
        'password': password,
    }
    try:
        response = session.post(login_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            token = result['data']
            headers['Authorization'] = token
            session.headers.update({'Authorization': token})
            return True, "欢迎回来！"
        else:
            return False, f"登录失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_user_info():
    user_info_url = f"{base_url}/userInfo"
    try:
        response = session.post(user_info_url, headers=headers, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            user_info = result['data']
            return True, user_info
        else:
            return False, f"获取用户信息失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_my_courses():
    my_courses_url = f"{base_url}/myCourses"
    try:
        response = session.post(my_courses_url, headers=headers, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            courses = result['data']
            return True, courses
        else:
            return False, f"获取课程列表失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_course_works(course_id):
    my_course_works_url = f"{base_url}/myCourseWorks"
    data = {
        'courseId': course_id,
    }
    try:
        response = session.post(my_course_works_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            works = result['data']
            return True, works
        else:
            return False, f"获取课程作业失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_questions(work_id):
    show_questions_url = f"{base_url}/showQuestions"
    data = {
        'workId': work_id,
    }
    try:
        response = session.post(show_questions_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            questions = result['data']
            return True, questions
        else:
            return False, f"获取题目失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def submit_answer(work_id, grade):
    submit_answer_url = f"{base_url}/submitAnswer"
    data = {
        'grade': grade,
        'workId': work_id,
    }
    try:
        response = session.post(submit_answer_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            return True, f"答案提交成功，成绩：{grade}\n返回信息：{result['data']}"
        else:
            return False, f"答案提交失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("理工学堂")
        self.root.resizable(False, False)

        self.style = ttk.Style("darkly")

        default_font = font.nametofont("TkDefaultFont")
        default_font.configure(family="微软雅黑", size=12)
        self.root.option_add("*Font", default_font)

        self.username = ""
        self.password = ""
        self.courses = []
        self.works = []
        self.questions = []
        self.images = []
        self.log_text = None
        self.selected_course_id = None
        self.selected_work_id = None
        self.selected_work_name = ""
        self.login_message = ""

        self.create_login_page()

    def create_login_page(self):
        self.clear_window()

        frame = ttk.Frame(self.root, padding=20)
        frame.pack(expand=True, fill=BOTH)

        ttk.Label(frame, text="欢迎使用理工学堂助手！", font=("微软雅黑", 14)).pack(pady=10)

        ttk.Label(frame, text="登录", font=("微软雅黑", 24, "bold")).pack(pady=20)

        username_frame = ttk.Frame(frame)
        username_frame.pack(fill=X, expand=True, pady=10)
        ttk.Label(username_frame, text="用户名：", width=10).pack(side=LEFT)
        self.username_entry = ttk.Entry(username_frame)
        self.username_entry.pack(side=LEFT, expand=True, fill=X)

        password_frame = ttk.Frame(frame)
        password_frame.pack(fill=X, expand=True, pady=10)
        ttk.Label(password_frame, text="密码：", width=10).pack(side=LEFT)
        self.password_entry = ttk.Entry(password_frame, show="*")
        self.password_entry.pack(side=LEFT, expand=True, fill=X)

        # 检查是否有已保存的用户名和密码
        saved_username = self.get_saved_username()
        saved_password = None
        if saved_username:
            self.username_entry.insert(0, saved_username)
            saved_password = self.get_saved_password(saved_username)
            if saved_password:
                self.password_entry.insert(0, saved_password)

        self.remember_var = tk.BooleanVar()
        if saved_username and saved_password:
            self.remember_var.set(True)
        else:
            self.remember_var.set(False)

        self.remember_check = ttk.Checkbutton(frame, text="记住密码", variable=self.remember_var)
        self.remember_check.pack(pady=5)

        ttk.Button(frame, text="登录", command=self.login, style='primary.TButton').pack(pady=20)

    def login(self):
        self.username = self.username_entry.get()
        self.password = self.password_entry.get()
        success, msg = login(self.username, self.password)
        if success:
            if self.remember_var.get():
                self.save_credentials(self.username, self.password)
            else:
                self.delete_saved_credentials()
            self.login_message = msg  # 存储登录成功的消息
            self.create_main_menu()
        else:
            Messagebox.show_error(message=msg, title="错误")

    @staticmethod
    def save_credentials(username, password):
        keyring.set_password(SERVICE_NAME, 'username', username)
        keyring.set_password(SERVICE_NAME, username, password)

    @staticmethod
    def get_saved_username():
        return keyring.get_password(SERVICE_NAME, 'username')

    @staticmethod
    def get_saved_password(username):
        return keyring.get_password(SERVICE_NAME, username)

    def delete_saved_credentials(self):
        saved_username = self.get_saved_username()
        if saved_username:
            keyring.delete_password(SERVICE_NAME, saved_username)
        keyring.delete_password(SERVICE_NAME, 'username')

    def create_main_menu(self):
        self.clear_window()
        frame = ttk.Frame(self.root, padding=20)
        frame.pack(expand=True, fill=BOTH)

        ttk.Label(frame, text="理工学堂助手", font=("微软雅黑", 24, "bold")).pack(pady=10)
        ttk.Label(frame, text="本软件免费使用，欢迎分享给朋友，但请勿倒卖或用于商业用途", font=("微软雅黑", 12)).pack(pady=5)

        if self.login_message:
            login_label = ttk.Label(frame, text=self.login_message, font=("微软雅黑", 12))
            login_label.pack(pady=10)
            self.root.after(3000, login_label.destroy)

        ttk.Button(frame, text="查看课程", command=self.view_courses, style='info.TButton').pack(pady=10, fill=X,
                                                                                                 expand=True)
        ttk.Button(frame, text="查看用户信息", command=self.view_user_info, style='info.TButton').pack(pady=10, fill=X,
                                                                                                       expand=True)
        ttk.Button(frame, text="退出", command=self.root.quit, style='danger.TButton').pack(pady=10, fill=X,
                                                                                            expand=True)

    @staticmethod
    def view_user_info():
        success, result = get_user_info()
        if success:
            info = result
            message = f"姓名：{info.get('name', 'N/A')}\n邮箱：{info.get('email', 'N/A')}"
            Messagebox.show_info(title="用户信息", message=message)
        else:
            Messagebox.show_error(title="错误", message=result)

    def view_courses(self):
        success, result = get_my_courses()
        if success:
            self.courses = result
            self.create_courses_page()
        else:
            Messagebox.show_error(title="错误", message=result)

    def create_courses_page(self):
        self.clear_window()
        frame = ttk.Frame(self.root, padding=20)
        frame.pack(expand=True, fill=BOTH)

        ttk.Label(frame, text="课程列表", font=("微软雅黑", 24, "bold")).pack(pady=10)
        courses_frame = ttk.Frame(frame)
        courses_frame.pack(pady=10, fill=BOTH, expand=True)

        for idx, course in enumerate(self.courses):
            btn = ttk.Button(courses_frame, text=f"{course['courseName']}",
                             command=lambda c_idx=idx: self.select_course(c_idx), style='success.TButton')
            btn.pack(pady=5, fill=X, expand=True)
        ttk.Button(frame, text="返回主菜单", command=self.create_main_menu, style='warning.TButton').pack(pady=20)
        ttk.Button(frame, text="全都提交100昏！！！", command=self.submit_all_courses_100,
                   style='warning.TButton').pack(pady=10)

    def select_course(self, idx):
        idx = int(idx)
        self.selected_course_id = self.courses[idx]['courseId']
        self.selected_course_name = self.courses[idx]['courseName']  # 保存课程名称
        success, result = get_course_works(self.selected_course_id)
        if success:
            self.works = result
            self.create_works_page()
        else:
            Messagebox.show_error(title="错误", message=result)

    def create_works_page(self):
        self.clear_window()
        frame = ttk.Frame(self.root, padding=20)
        frame.pack(expand=True, fill=BOTH)

        ttk.Label(frame, text="作业列表", font=("微软雅黑", 24, "bold")).pack(pady=10)
        works_frame = ttk.Frame(frame)
        works_frame.pack(pady=10, fill=BOTH, expand=True)

        if not self.works:
            ttk.Label(frame, text="该课程目前没有可用的作业。").pack(pady=20)
        else:
            for idx, work in enumerate(self.works):
                work_name = work['workName']
                work_id = work['workId']
                work_frame = ttk.Frame(works_frame)
                work_frame.pack(pady=5, fill=X, expand=True)

                ttk.Label(work_frame, text=work_name).pack(side=LEFT, padx=5)

                ttk.Button(work_frame, text="查看",
                           command=lambda w_idx=idx: self.select_work(w_idx), style='info.TButton').pack(side=RIGHT)

                ttk.Button(work_frame, text="收集所有题目",
                           command=lambda w_id=work_id, w_name=work_name: self.start_collecting_questions(w_id, w_name,
                                                                                                          self.selected_course_name),
                           style='primary.TButton').pack(side=RIGHT, padx=5)
        # 添加一键提交所有作业100分的按钮
        ttk.Button(frame, text="一键提交所有作业100分", command=self.submit_all_works_100,
                   style='warning.TButton').pack(pady=10)
        ttk.Button(frame, text="返回课程列表", command=self.create_courses_page, style='warning.TButton').pack(pady=20)

    # 一键提交所有作业100分
    def submit_all_works_100(self):
        # 使用线程来处理提交过程
        threading.Thread(target=self._submit_all_works_100).start()

    def _submit_all_works_100(self):
        success_count = 0
        fail_count = 0
        messages = []
        for work in self.works:
            work_id = work['workId']
            work_name = work['workName']
            success, result = submit_answer(work_id, '100')
            if success:
                success_count += 1
                messages.append(f"作业 '{work_name}' 提交成功。")
            else:
                fail_count += 1
                messages.append(f"作业 '{work_name}' 提交失败：{result}")
        # 在主线程中显示结果
        self.root.after(0, self._show_submit_results, success_count, fail_count, messages)

    @staticmethod
    def _show_submit_results(success_count, fail_count, messages):
        message = "\n".join(messages)
        message = f"成功提交 {success_count} 个作业，失败 {fail_count} 个作业。\n\n{message}"
        Messagebox.show_info(title="提交结果", message=message)

    def submit_all_courses_100(self):
        # 使用线程来处理提交过程
        threading.Thread(target=self._submit_all_courses_100).start()

    def _submit_all_courses_100(self):
        total_courses = len(self.courses)
        total_works = 0
        success_count = 0
        fail_count = 0
        messages = []
        for course in self.courses:
            course_id = course['courseId']
            course_name = course['courseName']
            success, result = get_course_works(course_id)
            if success:
                works = result
                total_works += len(works)
                for work in works:
                    work_id = work['workId']
                    work_name = work['workName']
                    submit_success, submit_result = submit_answer(work_id, '100')
                    if submit_success:
                        success_count += 1
                        messages.append(f"课程 '{course_name}' 的作业 '{work_name}' 提交成功。")
                    else:
                        fail_count += 1
                        messages.append(f"课程 '{course_name}' 的作业 '{work_name}' 提交失败：{submit_result}")
            else:
                messages.append(f"获取课程 '{course_name}' 的作业失败：{result}")
        # 在主线程中显示结果
        self.root.after(0, self._show_submit_all_courses_results, total_courses, total_works, success_count, fail_count,
                        messages)

    @staticmethod
    def _show_submit_all_courses_results(total_courses, total_works, success_count, fail_count, messages):
        message = "\n".join(messages)
        message = f"共处理 {total_courses} 门课程，{total_works} 个作业。\n成功提交 {success_count} 个作业，失败 {fail_count} 个作业。\n\n{message}"
        Messagebox.show_info(title="提交结果", message=message)

    # 启动收集题目
    def start_collecting_questions(self, work_id, work_name, course_name):
        assignment_name = ''.join(c for c in work_name if c not in r'<>:"/\|?*')
        course_name = ''.join(c for c in course_name if c not in r'<>:"/\|?*')

        assignment_name = f"{assignment_name} (ID_{work_id})"
        course_name = f"{course_name} (ID_{self.selected_course_id})"

        progress_window = tk.Toplevel(self.root)
        progress_window.title(f"收集作业 {work_name} 的题目")

        # 状态标签
        self.status_label = ttk.Label(progress_window, text="开始收集题目，请稍候...")
        self.status_label.pack(padx=10, pady=10)

        # 进度条
        self.progress_bar = ttk.Progressbar(progress_window, mode='indeterminate')
        self.progress_bar.pack(padx=10, pady=10, fill='x')
        self.progress_bar.start()

        # 使用线程避免阻塞GUI，传递课程名称
        threading.Thread(target=self.collect_all_questions,
                         args=(work_id, assignment_name, course_name, 100, progress_window)).start()

    # 收集题目方法
    def collect_all_questions(self, work_id, assignment_name, course_name, max_iterations=100, progress_window=None):
        assignment_folder = os.path.join('作业', course_name, assignment_name)
        if not os.path.exists(assignment_folder):
            os.makedirs(assignment_folder)

        collected_questions = {}
        no_new_questions_count = 0

        for i in range(max_iterations):
            success, questions = get_questions(work_id)
            if success:
                new_question_found = False
                for question in questions:
                    question_id = question.get('id')
                    if question_id not in collected_questions:
                        collected_questions[question_id] = question
                        # 保存题目和答案
                        self.save_question(question, assignment_folder)
                        new_question_found = True

                if new_question_found:
                    no_new_questions_count = 0
                    status_message = f"已收集到 {len(collected_questions)} 道题目。"
                else:
                    no_new_questions_count += 1
                    status_message = f"未发现新题目，已连续 {no_new_questions_count} 次未发现新题目。（最多重试10次）"
                    if no_new_questions_count >= 10:
                        status_message += "\n连续10次未获取到新题目，停止收集。"
                        # 更新状态标签
                        self.root.after(0, self.status_label.config, {'text': status_message})
                        break

                # 更新状态标签
                self.root.after(0, self.status_label.config, {'text': status_message})
                time.sleep(0.5)
            else:
                status_message = f"获取题目失败：{questions}"
                # 更新状态标签
                self.root.after(0, self.status_label.config, {'text': status_message})
                break

        # 保存所有题目信息到 Word 文档
        self.save_questions_to_word(collected_questions, assignment_folder)

        # 更新状态标签，停止进度条
        self.root.after(0, self.status_label.config, {'text': f"收集完成，共收集到 {len(collected_questions)} 道题目。"})
        self.root.after(0, self.progress_bar.stop)

        self.root.after(2000, progress_window.destroy)

    @staticmethod
    def save_questions_to_word(collected_questions, assignment_folder):
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        document.add_heading('题目和答案', 0)

        images_folder = os.path.join(assignment_folder, '题目图片')

        sorted_question_ids = sorted(collected_questions.keys(), key=lambda x: int(x))

        for idx, question_id in enumerate(sorted_question_ids):
            question = collected_questions[question_id]
            name = question.get('name', 'N/A')
            answer = question.get('answer', 'N/A')

            document.add_heading(f'题目 {idx + 1}', level=1)
            document.add_paragraph(f"题目ID：{question_id}")
            document.add_paragraph(f"题目名称：{name}")

            # 添加图片
            image_path = os.path.join(images_folder, f"{question_id}.png")
            if os.path.exists(image_path):
                document.add_picture(image_path, width=Inches(5))
            else:
                document.add_paragraph("（无图片）")

            document.add_paragraph(f"答案：{answer}")

        # 保存文档
        docx_path = os.path.join(assignment_folder, '题目和答案.docx')
        document.save(docx_path)

    # 保存题目方法
    def save_question(self, question, assignment_folder):
        question_id = question.get('id', 'N/A')
        imgurl = question.get('imgurl', 'N/A')

        # 创建保存图片的文件夹
        images_folder = os.path.join(assignment_folder, '题目图片')
        if not os.path.exists(images_folder):
            os.makedirs(images_folder)

        # 下载并保存图片
        if imgurl and imgurl != 'N/A':
            try:
                response = session.get(imgurl)
                response.raise_for_status()
                image_data = response.content
                image_path = os.path.join(images_folder, f"{question_id}.png")
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
            except Exception as e:
                self.log_text.insert(END, f"无法下载题目 {question_id} 的图片：{e}\n")
                self.log_text.see(END)

    def select_work(self, idx):
        idx = int(idx)
        self.selected_work_id = self.works[idx]['workId']
        self.selected_work_name = self.works[idx]['workName']  # 保存作业名称
        success, result = get_questions(self.selected_work_id)
        if success:
            self.questions = result
            self.create_questions_page()
        else:
            Messagebox.show_error(title="错误", message=result)

    def create_questions_page(self):
        self.clear_window()
        frame = ttk.Frame(self.root, padding=20)
        frame.pack(expand=True, fill=BOTH)

        ttk.Label(frame, text="题目列表", font=("微软雅黑", 24, "bold")).pack(pady=10)

        canvas_frame = ttk.Frame(frame)
        canvas_frame.pack(fill=BOTH, expand=True)

        canvas = tk.Canvas(canvas_frame)
        canvas.pack(side=tk.LEFT, fill=BOTH, expand=True)

        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        scrollbar.pack(side=tk.RIGHT, fill=Y)

        canvas.configure(yscrollcommand=scrollbar.set)

        questions_frame = ttk.Frame(canvas)
        canvas.create_window((0, 0), window=questions_frame, anchor='nw')

        def _on_mousewheel(event):
            canvas.yview_scroll(-int(event.delta / 120), "units")

        # 绑定鼠标滚轮事件
        if self.root.tk.call('tk', 'windowingsystem') == 'win32':
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
        elif self.root.tk.call('tk', 'windowingsystem') == 'x11':
            canvas.bind_all("<Button-4>", lambda event: canvas.yview_scroll(-1, 'units'))
            canvas.bind_all("<Button-5>", lambda event: canvas.yview_scroll(1, 'units'))
        elif self.root.tk.call('tk', 'windowingsystem') == 'aqua':
            canvas.bind_all("<MouseWheel>", _on_mousewheel)

        self.images = []  # 重置图像列表

        for idx, question in enumerate(self.questions):
            question_id = question.get('id', 'N/A')
            name = question.get('name', 'N/A')
            imgurl = question.get('imgurl', 'N/A')
            answer = question.get('answer', 'N/A')

            ttk.Label(questions_frame, text=f"{idx + 1}. 题目ID：{question_id}", font=("微软雅黑", 14, "bold")).pack(
                anchor='w', pady=5)
            ttk.Label(questions_frame, text=f"题目名称：{name}").pack(anchor='w')

            # 获取并显示图像
            if imgurl and imgurl != 'N/A':
                try:
                    response = session.get(imgurl)
                    response.raise_for_status()
                    image_data = response.content
                    image = Image.open(io.BytesIO(image_data))

                    # 获取原始尺寸
                    original_width, original_height = image.size

                    # 定义最大尺寸
                    max_width = 400
                    max_height = 300

                    # 计算缩放比例，保持宽高比
                    ratio = min(max_width / original_width, max_height / original_height)
                    if ratio < 1:
                        new_width = int(original_width * ratio)
                        new_height = int(original_height * ratio)
                        image = image.resize((new_width, new_height), Image.LANCZOS)
                    else:
                        _, new_height = original_width, original_height

                    photo = ImageTk.PhotoImage(image)
                    self.images.append(photo)
                    image_label = ttk.Label(questions_frame, image=photo)
                    image_label.pack(anchor='w', pady=5)
                except Exception:
                    ttk.Label(questions_frame, text="无法加载图片", foreground='red').pack(anchor='w')
            else:
                ttk.Label(questions_frame, text="没有图片").pack(anchor='w')

            ttk.Label(questions_frame, text=f"答案：{answer}").pack(anchor='w', pady=(0, 10))

        # 更新滚动区域
        questions_frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))

        grade_frame = ttk.Frame(frame)
        grade_frame.pack(pady=10)
        ttk.Label(grade_frame, text="提交成绩（0-100）：").pack(side='left', padx=5)
        self.grade_entry = ttk.Entry(grade_frame, width=10)
        self.grade_entry.pack(side='left', padx=5)
        ttk.Button(frame, text="提交", command=self.submit_grade, style='primary.TButton').pack(pady=10)
        ttk.Button(frame, text="返回作业列表", command=self.create_works_page, style='warning.TButton').pack(pady=10)

    def submit_grade(self):
        grade = self.grade_entry.get()
        if grade.isdigit() and 0 <= int(grade) <= 100:
            success, result = submit_answer(self.selected_work_id, grade)
            if success:
                Messagebox.show_info(title="提示", message=result)
            else:
                Messagebox.show_error(title="错误", message=result)
        else:
            Messagebox.show_error(title="错误", message="请输入有效的成绩（0-100）。")

    def clear_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()


if __name__ == "__main__":
    root = ttk.Window()
    app = App(root)
    root.mainloop()
